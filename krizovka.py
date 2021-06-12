from docx import oxml
from docx import Document
from docx.enum.style import WD_STYLE_TYPE

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from docx.shared import Cm

class CrosswordRow:

    def __init__(self, word, highlight):
        self.word = word
        self.highlight = highlight
        self.length = len(word)


class Crossword:

    def __init__(self):
        self.rows = []
        self.max_left = 0
        self.max_right = 0
        self.row_num = 0
        self.columns = 0

    def addRow(self, row):
        self.rows.append(row)
        h_index = row.highlight
        if self.max_left < h_index:
            self.max_left = h_index
        if self.max_right < row.length - h_index:
            self.max_right = row.length - h_index - 1
        self.row_num += 1
        self.columns = self.max_left + self.max_right + 2

    def getColNum(self):
        return 2 + self.max_left + self.max_right

    def __str__(self):
        for word in self.rows:
            print(word.word)


def main():
    words = str(input("Zadej slova do křížovky (seřazené od shora dolů, oddělené mezerou)\n"))
    secret = str(input("Zadej tajenku\n"))
    crossword = init_crossword(words, secret)
    print_crossword(crossword)


def init_crossword(words, secret):
    rows = words.split(",")
    secret = secret.strip()
    crossword = Crossword()

    if len(rows) != len(secret):
        exit(1)
    for i in range(len(rows)):
        h_letter = secret[i]
        row = rows[i]
        row = row.strip()
        if row == "" and secret[i] == " ":
            c_row = CrosswordRow(" ", 0)
            crossword.addRow(c_row)
            continue
        indices = []
        for j, r_letter in enumerate(row):
            if r_letter == h_letter:
                indices.append(j)
        if not indices:
            print(f"slovo {row} neobsahuje písmeno z tajenky {h_letter.capitalize()}")

        c_row = CrosswordRow(row, indices[len(indices) // 2])
        crossword.addRow(c_row)

    return crossword


def set_style(styles):
    style = styles.add_style('Table letter', WD_STYLE_TYPE.PARAGRAPH)
    style.font.bold = True
    style.font.size = Pt(16)
    style.font.name = 'Calibri'


def print_crossword(crossword):
    document = Document()

    table_vis = document.add_table(rows=crossword.row_num, cols=crossword.getColNum(), style='TableGrid')
    document.add_paragraph()
    table_invis = document.add_table(rows=crossword.row_num, cols=crossword.getColNum(), style='TableGrid')
    l_max = crossword.max_left

    set_style(document.styles)

    format_table(table_vis, crossword)
    format_table(table_invis, crossword)

    for i, row in enumerate(crossword.rows):
        print_cross_row(row, i, l_max, True, table_vis)
        print_cross_row(row, i, l_max, False, table_invis)


    document.save("krizovka.docx")


def print_cross_row(row, index, max_left, visible, table):
    table_columns = table.rows[index].cells
    table_columns[0].text = str(index + 1)
    table_columns[0].paragraphs[0].style = "Table letter"
    highlight = row.highlight
    j = 0

    for i in range(1 + max_left - highlight, 1 + max_left - highlight + row.length):
        if visible:
            table_columns[i].text = row.word[j].capitalize()

        if j == highlight:
            shading_red_1 = oxml.parse_xml(r'<w:shd {} w:fill="FFFF66"/>'.
                                                format(oxml.ns.nsdecls('w')))
            table_columns[i]._tc.get_or_add_tcPr().append(shading_red_1)
        table_columns[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        table_columns[i].paragraphs[0].style = "Table letter"
        table_columns[i].width = Cm(1)
        j += 1


def format_table(table, crossword):
    tbl = table._tbl  # get xml element in table
    for row in table.rows:
        row.height = Cm(1)

    row_num = 0
    cell_num = 0
    row_highlight = crossword.rows[row_num].highlight

    row_left_index = 1 + crossword.max_left - row_highlight
    row_right_index = row_left_index + crossword.rows[row_num].length - 1

    print(crossword.columns)
    for cell in tbl.iter_tcs():

        cell.width = Cm(1)
        if cell_num == crossword.columns:
            cell_num = 0
            row_num += 1
            row_highlight = crossword.rows[row_num].highlight

            row_left_index = 1 + crossword.max_left - row_highlight
            row_right_index = row_left_index + crossword.rows[row_num].length - 1

        if row_left_index <= cell_num <= row_right_index:

            cell_num += 1
            continue


        tcPr = cell.tcPr  # get tcPr element, in which we can define style of borders
        tcBorders = OxmlElement('w:tcBorders')
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'nil')

        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'nil')

        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'nil')

        right = OxmlElement('w:right')
        right.set(qn('w:val'), 'nil')

        tcBorders.append(top)
        tcBorders.append(left)
        tcBorders.append(bottom)
        tcBorders.append(right)
        tcPr.append(tcBorders)

        cell_num += 1



main()



